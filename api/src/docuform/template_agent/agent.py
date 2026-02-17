"""
Template Field Detection Agent

An AI agent that analyzes DOCX templates and helps attorneys identify and create
fillable fields. The agent can detect placeholders, suggest field names, and
convert text into template fields.
"""
from __future__ import annotations

import logfire
from pathlib import Path
from dataclasses import dataclass, field
from docx import Document
from pydantic import BaseModel
from pydantic_ai import Agent, RunContext
from pydantic_ai.models.openai import OpenAIChatModel

from api.src.docuform.docx_content_controls import (
    wrap_text_in_content_control,
    read_content_controls_detailed,
    set_content_control_value,
    delete_content_control,
    update_content_control,
    split_content_control,
    W_SDT,
    W_SDT_PR,
    W_SDT_CONTENT,
    W_TAG,
    W_VAL,
    W_T,
    W_R,
)


# Directory for documents
DOCUMENTS_DIR = Path(__file__).parent.parent / "documents"

# Working copy retention (in seconds) - cleanup files older than this
WORKING_COPY_MAX_AGE_SECONDS = 3600  # 1 hour


def cleanup_old_working_copies() -> int:
    """
    Clean up old working copy files on module import.
    Removes working copies older than WORKING_COPY_MAX_AGE_SECONDS.
    Returns the number of files deleted.
    """
    import time
    deleted = 0
    now = time.time()

    # Match pattern: *_template_agent_working_*.docx (conversation-scoped working copies)
    for file in DOCUMENTS_DIR.glob("*_template_agent_working_*.docx"):
        try:
            age = now - file.stat().st_mtime
            if age > WORKING_COPY_MAX_AGE_SECONDS:
                file.unlink()
                deleted += 1
                logfire.info(f"Cleaned up old working copy: {file.name}", age_seconds=age)
        except Exception as e:
            logfire.warn(f"Failed to clean up {file.name}: {e}")

    if deleted > 0:
        logfire.info(f"Cleaned up {deleted} old working copies")

    return deleted


# Run cleanup on module import
cleanup_old_working_copies()


class DetectedField(BaseModel):
    """A field detected in the document"""
    text: str
    suggested_tag: str
    suggested_alias: str
    reason: str


class FieldDetectionResult(BaseModel):
    """Result of field detection analysis"""
    fields: list[DetectedField]
    summary: str


class FieldToCreate(BaseModel):
    """Input for creating a template field from text"""
    search_text: str  # The exact text to find and convert to a field
    tag: str  # The programmatic tag/key for the field (dot notation, e.g., "declarant.name")
    alias: str | None = None  # Human-readable display name (defaults to formatted tag)
    replace_with_placeholder: bool = True  # If True, replace text with "[tag]" placeholder


@dataclass
class TemplateAgentContext:
    """Context for the template agent, holding the working document state"""
    document_filename: str  # Required - the document to work with
    conversation_id: str | None = None  # Conversation ID for scoped working copies
    document: Document | None = field(default=None, repr=False)
    working_filename: str | None = None  # Temp file for modifications

    def get_document_path(self) -> Path:
        """Get the path to the current document"""
        return DOCUMENTS_DIR / self.document_filename

    def get_working_path(self) -> Path | None:
        """Get the path to the working copy (scoped by conversation_id if available)"""
        if self.working_filename:
            return DOCUMENTS_DIR / self.working_filename
        return None

    def load_document(self) -> tuple[bool, str]:
        """Load the document from filename. Returns (success, message).

        If a working copy exists for this conversation, loads from that to preserve
        previous modifications. Otherwise loads from the original document.
        Working copies are scoped by conversation_id to isolate multi-user scenarios.
        """
        original_path = self.get_document_path()
        if not original_path.exists():
            available = [f.name for f in DOCUMENTS_DIR.glob("*.docx")]
            return False, f"Document '{self.document_filename}' not found. Available: {available}"

        # Set working filename - scoped by conversation_id
        # Format: {stem}_template_agent_working_{short_id}.docx
        if self.conversation_id:
            # Use first 8 chars of conversation_id for shorter filenames
            short_id = self.conversation_id[:8]
            self.working_filename = f"{original_path.stem}_template_agent_working_{short_id}.docx"
        else:
            # Fallback for backwards compatibility (shouldn't happen in practice)
            self.working_filename = f"{original_path.stem}_template_agent_working.docx"

        working_path = self.get_working_path()

        try:
            # Load from working copy if it exists (preserves previous modifications)
            if working_path and working_path.exists():
                self.document = Document(str(working_path))
                return True, f"Loaded {self.document_filename} (with previous modifications)"
            else:
                self.document = Document(str(original_path))
                return True, f"Loaded {self.document_filename}"
        except Exception as e:
            return False, f"Error loading document: {str(e)}"


# Create the agent
model = OpenAIChatModel("gpt-4o")

agent = Agent(
    name="template_agent",
    model=model,
    deps_type=TemplateAgentContext,
    instructions="""You are a legal document template assistant. You help attorneys convert documents
into templates with fillable fields.

IMPORTANT: A document has been pre-selected and loaded for this session. At the START of EVERY conversation,
call `get_document_info` to see which document is loaded.

## Document Types You'll Encounter

1. **Existing templates** - Documents with placeholders like [Name], {{field}}, or blank lines.
   These just need the placeholders marked as template fields.

2. **Filled documents** - Real documents with actual names, dates, addresses, amounts.
   The user wants to CONVERT these into templates by replacing specific values with placeholders.
   Example: "Allen R. Moreland" → [declarant.name], "January 15, 2024" → [signing.date]

If you see a document with real names, dates, and addresses (not placeholders), ASK the user:
"This looks like a filled document rather than a template. Would you like me to help convert it
into a template by identifying values that should become fillable fields?"

## Your Tools

**Reading & Searching:**
- `get_document_info` - See loaded document info and preview
- `get_document_text` - Get the full document text for analysis
- `list_fields` - Show existing template fields
- `find_text` - Search or browse all document text (body, tables, headers, footers)

**Modifying:**
- `create_fields` - Convert text into template fields. Always use this tool for creating fields.
- `edit_field` - Edit an existing field's tag, display name, or value
- `delete_field` - Remove a field (optionally preserving its text)
- `split_field` - Split one field into multiple adjacent fields (e.g., name → first/middle/last)
- `update_field_value` - Quick way to change just the value of an existing field
- `reset_working_copy` - Discard all modifications

## Workflow for Converting Filled Documents

1. Use `get_document_text` to read the full content
2. Identify values that should become fields (names, dates, addresses, amounts)
3. Use `find_text` to find all occurrences of each value
4. Ask user to confirm which items to convert
5. Use `create_fields` to mark ALL confirmed fields in one call
6. Run `get_document_text` again to see if there are any remaining instances of the same/similar text, or if was an already filled document being converted into a template there could be straggler fields that need to be created.

## Tag Naming
ALWAYS use dot notation (object.property format) for clarity and consistency:

**People/Entities:**
- `declarant.name`, `declarant.address`, `declarant.date_of_birth`
- `spouse.name`, `spouse.address`
- `child1.name`, `child1.date_of_birth`, `child2.name`, `child2.date_of_birth`
- `witness1.name`, `witness1.address`, `witness2.name`
- `executor.name`, `executor.address`
- `beneficiary1.name`, `beneficiary1.share_percentage`
- `guardian.name`, `guardian.address`

**Dates/Events:**
- `signing.date`, `signing.state`, `signing.city`
- `document.effective_date`, `document.expiration_date`
- `closing.date`, `closing.location`

**Transactions/Amounts:**
- `purchase.price`, `purchase.deposit`
- `lease.start_date`, `lease.end_date`, `lease.monthly_rent`

Display names should be human-readable:
- `declarant.name` → "Declarant Name"
- `child1.date_of_birth` → "Child1 Date Of Birth"
- `signing.date` → "Signing Date"

## When Field Creation Fails
If create_fields can't find text, PROACTIVELY diagnose using `find_text`:

1. **Search mode**: `find_text(query="the text")` searches the ENTIRE document
   (body, tables, headers, footers). Text is often in tables, especially signature blocks.

2. **Filter by location**: `find_text(query="text", location="table")` to search
   only in tables, or use location="body", "header", "footer".

3. **Browse mode**: `find_text()` with no query lists all segments. Use pagination
   with `start` and `limit` params, or filter with `location="table"`.

4. **Try shorter substrings** - search for just a first name or part of a word.

NOTE: The field creation tool handles text that spans multiple document segments automatically.
You don't need to split names like "John Smith" into separate fields - just search for the
full text and create_fields will combine the segments into one field.

IMPORTANT: Do this investigation automatically without asking the user. Only communicate
results in user-friendly terms like "I found it in a table".

## Splitting a Field into Multiple Fields
Use the `split_field` tool to replace one field with multiple adjacent fields:
- `split_field(tag="declarant.name", new_tags=["declarant.first_name", "declarant.middle_name", "declarant.last_name"])`

This replaces the single field with placeholder fields like `[declarant.first_name] [declarant.middle_name] [declarant.last_name]`.
The original text is discarded - users will fill in the new fields when using the template.

## Finishing Up
When the user is happy with the template and all fields have been created, tell them to click
the **Save** or **Save As** button in the interface to save their completed template.
This is the final step - remind them that the save buttons are how they get their finished document.

Always explain what you're doing and ask for confirmation before making changes.
""",
    retries=2,
)


@agent.instructions
async def add_document_context(ctx: RunContext[TemplateAgentContext]) -> str:
    """Dynamic system prompt that loads and describes the document."""
    print(f"🔥 document_context CALLED with filename: {ctx.deps.document_filename}")
    try:
        logfire.info(f"document_context called with filename: {ctx.deps.document_filename}")

        # Load the document if not already loaded
        if ctx.deps.document is None:
            logfire.info(f"Loading document: {ctx.deps.document_filename}")
            success, message = ctx.deps.load_document()
            logfire.info(f"Load result: success={success}, message={message}")
            if not success:
                return f"ERROR: {message}. Please inform the user."

        # Get document stats
        doc = ctx.deps.document
        para_count = len(doc.paragraphs) if doc else 0
        controls = read_content_controls_detailed(str(ctx.deps.get_document_path()))

        # Get text preview (first 500 chars)
        text_preview = ""
        if doc:
            all_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            text_preview = all_text[:500] + ("..." if len(all_text) > 500 else "")

        result = f"""
The document "{ctx.deps.document_filename}" is already loaded and ready for analysis.

Document Info:
- Paragraphs: {para_count}
- Existing template fields: {len(controls)}

Text Preview:
{text_preview}

Start by analyzing this document for potential fields, or ask the user what they'd like to do.
"""
        logfire.info(f"document_context returning prompt (len={len(result)})")
        return result
    except Exception as e:
        logfire.error(f"Error in document_context: {e}", exc_info=True)
        return f"ERROR loading document context: {e}"


@agent.tool
async def get_document_info(ctx: RunContext[TemplateAgentContext]) -> str:
    """
    Get information about the currently loaded document.

    Call this at the START of every conversation to see which document is loaded
    and get a preview of its contents.

    Returns:
        Document filename, stats, and text preview
    """
    if ctx.deps.document is None:
        return f"Document '{ctx.deps.document_filename}' is set but not loaded. This is unexpected - contact support."

    doc = ctx.deps.document
    para_count = len(doc.paragraphs)
    controls = read_content_controls_detailed(str(ctx.deps.get_document_path()))

    # Get text preview (first 1000 chars)
    all_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    text_preview = all_text[:1000] + ("..." if len(all_text) > 1000 else "")

    logfire.info(f"get_document_info called", filename=ctx.deps.document_filename, para_count=para_count)

    return f"""Document: {ctx.deps.document_filename}

Stats:
- Paragraphs: {para_count}
- Existing template fields: {len(controls)}

Text Preview:
{text_preview}

Existing template fields:
{_format_fields(controls) if controls else "None"}

You can now analyze this document for potential fields, or ask what the user wants to do."""


@agent.tool
async def load_document(ctx: RunContext[TemplateAgentContext], filename: str) -> str:
    """
    Load a DOCX document for analysis and modification.

    Args:
        filename: Name of the document file (e.g., "contract.docx")

    Returns:
        Status message with document info
    """
    file_path = DOCUMENTS_DIR / filename

    if not file_path.exists():
        available = [f.name for f in DOCUMENTS_DIR.glob("*.docx")]
        return f"Document '{filename}' not found. Available documents: {available}"

    try:
        doc = Document(str(file_path))
        ctx.deps.document = doc
        ctx.deps.document_filename = filename

        # Create working copy filename
        stem = file_path.stem
        ctx.deps.working_filename = f"{stem}_working.docx"

        # Count paragraphs and existing controls
        para_count = len(doc.paragraphs)
        controls = read_content_controls_detailed(str(file_path))

        # Get text preview
        text_preview = "\n".join([p.text for p in doc.paragraphs[:10] if p.text.strip()])
        if len(doc.paragraphs) > 10:
            text_preview += "\n..."

        logfire.info(f"Loaded document: {filename}", para_count=para_count, control_count=len(controls))

        return f"""Document loaded: {filename}
- Paragraphs: {para_count}
- Existing template fields: {len(controls)}

Preview:
{text_preview}

Existing template fields:
{_format_fields(controls) if controls else "None"}"""

    except Exception as e:
        logfire.error(f"Failed to load document: {e}")
        return f"Error loading document: {str(e)}"


@agent.tool
async def get_document_text(ctx: RunContext[TemplateAgentContext], max_chars: int = 10000) -> str:
    """
    Get the full text content of the document for analysis.

    Use this to read the document content and identify potential fields yourself.
    Look for things like names, dates, addresses, dollar amounts, or any text
    that should be replaceable in a template.

    Existing template fields are shown as {{FIELD:tag_name:value}} so you know
    what has already been processed.

    Args:
        max_chars: Maximum characters to return (default 10000). Use higher values for longer documents.

    Returns:
        The document text content with fields marked
    """
    if ctx.deps.document is None:
        return "No document loaded. Use load_document first."

    doc = ctx.deps.document

    # Build text from all runs including content controls
    all_runs = _collect_all_runs(doc)

    # Group runs by paragraph to reconstruct readable text
    paragraphs: dict[str, list[tuple[str, str]]] = {}
    for loc, text, _ in all_runs:
        # Extract paragraph key (everything before last colon segment)
        # e.g., "body:p0:r0" -> "body:p0", "body:p0:field[name]" -> "body:p0"
        parts = loc.rsplit(":", 1)
        para_key = parts[0] if len(parts) > 1 else loc

        if para_key not in paragraphs:
            paragraphs[para_key] = []

        # Mark fields clearly in the text stream
        if ":field[" in loc:
            # Extract tag name from location like "body:p0:field[declarant.name]"
            tag_start = loc.find(":field[") + 7
            tag_end = loc.find("]", tag_start)
            tag_name = loc[tag_start:tag_end]
            paragraphs[para_key].append((loc, f"{{{{FIELD:{tag_name}:{text}}}}}"))
        else:
            paragraphs[para_key].append((loc, text))

    # Join runs within each paragraph, then join paragraphs
    para_texts = []
    for para_key in sorted(paragraphs.keys()):
        para_text = "".join(text for _, text in paragraphs[para_key])
        if para_text.strip():
            para_texts.append(para_text)

    all_text = "\n\n".join(para_texts)

    if len(all_text) > max_chars:
        return all_text[:max_chars] + f"\n\n... (truncated, {len(all_text) - max_chars} more characters)"

    return all_text


@agent.tool
async def list_fields(ctx: RunContext[TemplateAgentContext]) -> str:
    """
    List all existing template fields in the document.

    Returns:
        List of fields with their tags, display names, and current values
    """
    if ctx.deps.document is None:
        return "No document loaded. Use load_document first."

    # Read from working copy if it exists, otherwise from original
    working_path = ctx.deps.get_working_path()
    if working_path and working_path.exists():
        controls = read_content_controls_detailed(str(working_path))
    else:
        controls = read_content_controls_detailed(str(ctx.deps.get_document_path()))

    if not controls:
        return "No template fields found in the document."

    result = f"Found {len(controls)} template field(s):\n\n"
    for i, c in enumerate(controls, 1):
        result += f"{i}. Tag: {c['tag']}\n"
        result += f"   Display Name: {c['alias']}\n"
        result += f"   Current Value: \"{c['value']}\"\n\n"

    return result


def _extract_paragraph_content(para_element, location_prefix: str) -> list[tuple[str, str, any]]:
    """
    Extract text from a paragraph element, including both regular runs and content controls.

    Iterates through paragraph children in document order to capture:
    - Regular runs (w:r elements)
    - Content controls (w:sdt elements) - marked with [field:tag] prefix

    Returns list of (location, text, element) tuples.
    """
    results: list[tuple[str, str, any]] = []
    run_idx = 0
    sdt_idx = 0

    for child in para_element:
        tag = child.tag

        if tag == W_R:
            # Regular run - extract text
            texts = [t.text for t in child.iter(W_T) if t.text]
            text = "".join(texts)
            if text:
                results.append((f"{location_prefix}:r{run_idx}", text, child))
            run_idx += 1

        elif tag == W_SDT:
            # Content control - extract text and tag name
            sdt_pr = child.find(W_SDT_PR)
            tag_el = sdt_pr.find(W_TAG) if sdt_pr is not None else None
            field_tag = tag_el.get(W_VAL) if tag_el is not None else "unknown"

            content = child.find(W_SDT_CONTENT)
            texts = [t.text for t in content.iter(W_T) if t.text] if content is not None else []
            text = "".join(texts)

            if text:
                # Mark as field with tag name for AI context
                results.append((f"{location_prefix}:field[{field_tag}]", text, child))
            sdt_idx += 1

    return results


def _collect_all_runs(doc: Document) -> list[tuple[str, str, any]]:
    """
    Collect ALL text from the entire document, including:
    - Body paragraphs (regular text and content controls)
    - Tables (all cells)
    - Headers and footers

    Content controls are marked with location like "body:p5:field[tag_name]"
    so the AI knows which text is inside template fields.

    Returns list of (location, text, element) tuples.
    """
    all_runs: list[tuple[str, str, any]] = []

    # Body paragraphs
    for para_idx, para in enumerate(doc.paragraphs):
        all_runs.extend(_extract_paragraph_content(para._element, f"body:p{para_idx}"))

    # Tables
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para_idx, para in enumerate(cell.paragraphs):
                    loc_prefix = f"table{table_idx}:row{row_idx}:cell{cell_idx}:p{para_idx}"
                    all_runs.extend(_extract_paragraph_content(para._element, loc_prefix))

    # Headers and footers (from all sections)
    for section_idx, section in enumerate(doc.sections):
        # Header
        try:
            header = section.header
            if header:
                for para_idx, para in enumerate(header.paragraphs):
                    all_runs.extend(_extract_paragraph_content(para._element, f"header{section_idx}:p{para_idx}"))
        except Exception:
            pass

        # Footer
        try:
            footer = section.footer
            if footer:
                for para_idx, para in enumerate(footer.paragraphs):
                    all_runs.extend(_extract_paragraph_content(para._element, f"footer{section_idx}:p{para_idx}"))
        except Exception:
            pass

    return all_runs


@agent.tool
async def find_text(
    ctx: RunContext[TemplateAgentContext],
    query: str | None = None,
    location: str | None = None,
    start: int = 0,
    limit: int = 30,
    context: int = 3,
) -> str:
    """
    Search or browse all document text including body, tables, headers, and footers.

    Use this to:
    - Find specific text that should become template fields
    - Check for remaining instances after creating fields (e.g., name in header and footer)
    - Locate text that might be in tables or headers

    Two modes:
    - Search mode (query provided): Find text and show surrounding context
    - Browse mode (no query): List all text segments with pagination

    Args:
        query: Text to search for (case-insensitive). None = browse all segments.
        location: Filter by location: "body", "table", "header", "footer" (optional)
        start: Pagination offset for browse mode (default: 0)
        limit: Max results/segments to return (default: 30, max: 100)
        context: Segments to show around matches in search mode (default: 3)

    Returns:
        Found text locations or list of document segments
    """
    if ctx.deps.document is None:
        return "No document loaded. Use load_document first."

    doc = ctx.deps.document
    all_runs = _collect_all_runs(doc)

    if not all_runs:
        return "Document has no text content."

    # Apply location filter
    if location:
        location_lower = location.lower()
        all_runs = [(loc, text, run) for loc, text, run in all_runs if location_lower in loc.lower()]
        if not all_runs:
            return f"No segments found in '{location}'"

    limit = min(limit, 100)

    def format_segment(loc: str, text: str) -> str:
        """Format a segment, marking fields clearly."""
        display_text = text[:80] + ("..." if len(text) > 80 else "")
        if ":field[" in loc:
            # Extract field tag from location like "body:p5:field[declarant.name]"
            return f"[{loc}]: (FIELD) \"{display_text}\""
        return f"[{loc}]: \"{display_text}\""

    # Browse mode: list segments with pagination
    if not query:
        end = min(start + limit, len(all_runs))
        result = f"Segments {start+1}-{end} of {len(all_runs)} total"
        if location:
            result += f" (in '{location}')"
        result += ":\n\n"

        for i in range(start, end):
            loc, text, _ = all_runs[i]
            result += f"{i+1}. {format_segment(loc, text)}\n"

        if end < len(all_runs):
            result += f"\n... {len(all_runs) - end} more. Use start={end} to continue."

        return result

    # Search mode: find text with context
    query_lower = query.lower()

    # Build combined text and map positions back to runs
    combined_text = ""
    run_boundaries: list[tuple[int, int, int]] = []

    for i, (loc, text, _) in enumerate(all_runs):
        seg_start = len(combined_text)
        combined_text += text
        seg_end = len(combined_text)
        run_boundaries.append((seg_start, seg_end, i))

    # Find all occurrences
    results = []
    pos = 0
    while len(results) < limit:
        idx = combined_text.lower().find(query_lower, pos)
        if idx == -1:
            break

        match_end = idx + len(query)

        # Find the run(s) containing the match
        matching_run_indices = []
        for seg_start, seg_end, run_idx in run_boundaries:
            if seg_start < match_end and seg_end > idx:
                matching_run_indices.append(run_idx)

        if matching_run_indices:
            first_idx = matching_run_indices[0]
            last_idx = matching_run_indices[-1]

            # Get context runs
            ctx_start = max(0, first_idx - context)
            ctx_end = min(len(all_runs), last_idx + context + 1)

            # Format the runs
            runs_info = []
            for i in range(ctx_start, ctx_end):
                loc, text, _ = all_runs[i]
                marker = ">>>" if i in matching_run_indices else "   "
                field_marker = " (FIELD)" if ":field[" in loc else ""
                runs_info.append(f"{marker} [{loc}]:{field_marker} \"{text}\"")

            actual_match = combined_text[idx:match_end]
            result = f"Match {len(results) + 1}: \"{actual_match}\"\n"
            result += "\n".join(runs_info)
            results.append(result)

        pos = idx + 1

    if not results:
        # Summarize where text exists
        locations = set()
        for loc, _, _ in all_runs:
            if loc.startswith("body:"):
                locations.add("body")
            elif loc.startswith("table"):
                locations.add("tables")
            elif loc.startswith("header"):
                locations.add("headers")
            elif loc.startswith("footer"):
                locations.add("footers")

        return (
            f"'{query}' not found.\n\n"
            f"Document has {len(all_runs)} segments in: {', '.join(sorted(locations))}\n"
            f"Try a shorter substring or browse with query=None."
        )

    header = f"Found {len(results)} match(es) for '{query}':\n"
    header += ">>> = matching segment. Location: [container:paragraph:run]\n\n"

    return header + "\n\n".join(results)


def _create_single_field(
    document: Document,
    search_text: str,
    tag: str,
    alias: str | None = None,
    replace_with_placeholder: bool = True,
) -> tuple[bool, str]:
    """
    Helper function to convert text into a template field (content control internally).

    Args:
        document: The python-docx Document object
        search_text: The exact text to find and convert
        tag: The programmatic tag/key for the field (dot notation, e.g., "declarant.name")
        alias: Human-readable display name (defaults to formatted tag if not provided)
        replace_with_placeholder: If True (default), replace the text with "[tag]" placeholder

    Returns:
        Tuple of (success: bool, message: str)
    """
    if alias is None:
        # Convert dot notation to human-readable: "child1.date_of_birth" → "Child1 Date Of Birth"
        alias = tag.replace(".", " ").replace("_", " ").title()

    try:
        # Internal: uses content controls to implement fields
        count = wrap_text_in_content_control(
            document,
            search_text=search_text,
            tag=tag,
            alias=alias,
            first_only=False
        )

        if count > 0:
            placeholder_text = f"[{tag}]"
            if replace_with_placeholder:
                set_content_control_value(document, tag, placeholder_text)
                return True, f"'{search_text}' → '{tag}' → '{placeholder_text}'"
            return True, f"'{search_text}' → '{tag}'"
        else:
            return False, f"Text not found: '{search_text}'"

    except Exception as e:
        return False, f"Error creating field from '{search_text}': {str(e)}"


@agent.tool(sequential=True)
async def create_fields(
    ctx: RunContext[TemplateAgentContext],
    fields: list[FieldToCreate],
) -> str:
    """
    Convert text into template fields. Use this tool to mark text as fillable fields.

    Args:
        fields: List of fields to create. Each field has:
            - search_text: The exact text to find and convert
            - tag: The programmatic tag/key (dot notation, e.g., "declarant.name")
            - alias: Optional human-readable name (defaults to formatted tag)
            - replace_with_placeholder: If True (default), replace text with "[tag]"

    Example:
        fields=[
            FieldToCreate(search_text="Allen R. Moreland", tag="declarant.name"),
            FieldToCreate(search_text="January 15, 2024", tag="signing.date"),
            FieldToCreate(search_text="$50,000", tag="purchase.price", alias="Purchase Price")
        ]

    Returns:
        Summary of field creation with successes and failures
    """
    if ctx.deps.document is None:
        return "No document loaded. Use load_document first."

    if not fields:
        return "No fields provided."

    results = []
    successes = 0
    failures = 0

    for field in fields:
        success, message = _create_single_field(
            document=ctx.deps.document,
            search_text=field.search_text,
            tag=field.tag,
            alias=field.alias,
            replace_with_placeholder=field.replace_with_placeholder,
        )

        if success:
            results.append(f"✓ {message}")
            successes += 1
        else:
            results.append(f"❌ {message}")
            failures += 1

    # Save working copy once after all operations
    try:
        working_path = ctx.deps.get_working_path()
        if working_path:
            ctx.deps.document.save(str(working_path))
    except Exception as e:
        return f"Operations completed but failed to save: {str(e)}\n\nResults:\n" + "\n".join(results)

    logfire.info(f"Batch create_fields completed", successes=successes, failures=failures)

    summary = f"Created {successes} field(s)"
    if failures > 0:
        summary += f", {failures} failed"
    summary += ":\n\n" + "\n".join(results)

    return summary


@agent.tool(sequential=True)
async def update_field_value(
    ctx: RunContext[TemplateAgentContext],
    tag: str,
    new_value: str,
) -> str:
    """
    Update the text value inside an existing template field.

    Use this to change the displayed text of a field - for example,
    replacing "Allen R. Moreland" with "[declarant.name]" as a placeholder.

    Args:
        tag: The tag of the field to modify (e.g., "declarant.name")
        new_value: The new text value to set (e.g., "[declarant.name]")

    Returns:
        Status message indicating success or failure
    """
    if ctx.deps.document is None:
        return "No document loaded. Use load_document first."

    try:
        # Get current value for logging
        working_path = ctx.deps.get_working_path()
        if working_path and working_path.exists():
            controls = read_content_controls_detailed(str(working_path))
        else:
            controls = read_content_controls_detailed(str(ctx.deps.get_document_path()))

        old_value = None
        for ctrl in controls:
            if ctrl["tag"] == tag:
                old_value = ctrl["value"]
                break

        if old_value is None:
            available_tags = [ctrl["tag"] for ctrl in controls]
            return f"Field with tag '{tag}' not found. Available tags: {available_tags}"

        # Replace the text
        success = set_content_control_value(ctx.deps.document, tag, new_value)

        if success:
            # Save working copy
            if working_path:
                ctx.deps.document.save(str(working_path))

            logfire.info(f"Updated field {tag}: '{old_value}' -> '{new_value}'")
            return f"Updated field '{tag}': '{old_value}' → '{new_value}'"
        else:
            return f"Failed to update field '{tag}'"

    except Exception as e:
        logfire.error(f"Failed to update field: {e}")
        return f"Error updating field: {str(e)}"


@agent.tool(sequential=True)
async def delete_field(
    ctx: RunContext[TemplateAgentContext],
    tag: str,
    preserve_text: bool = True,
) -> str:
    """
    Delete a template field from the document.

    Args:
        tag: The tag of the field to delete (e.g., "declarant.name")
        preserve_text: If True (default), keep the text content in place.
                       If False, remove both the field and its text.

    Returns:
        Status message indicating success or failure
    """
    if ctx.deps.document is None:
        return "No document loaded. Use load_document first."

    try:
        # Get current value for logging
        working_path = ctx.deps.get_working_path()
        if working_path and working_path.exists():
            controls = read_content_controls_detailed(str(working_path))
        else:
            controls = read_content_controls_detailed(str(ctx.deps.get_document_path()))

        # Check if field exists
        field_info = None
        for ctrl in controls:
            if ctrl["tag"] == tag:
                field_info = ctrl
                break

        if field_info is None:
            available_tags = [ctrl["tag"] for ctrl in controls]
            return f"Field with tag '{tag}' not found. Available tags: {available_tags}"

        # Delete the field
        count = delete_content_control(ctx.deps.document, tag, preserve_text=preserve_text)

        if count > 0:
            # Save working copy
            if working_path:
                ctx.deps.document.save(str(working_path))

            logfire.info(f"Deleted field {tag}", preserve_text=preserve_text)
            return f"Deleted field '{tag}' (was: '{field_info['value'][:50]}...')" if len(field_info['value']) > 50 else f"Deleted field '{tag}' (was: '{field_info['value']}')"
        else:
            return f"Failed to delete field '{tag}'"

    except Exception as e:
        logfire.error(f"Failed to delete field: {e}")
        return f"Error deleting field: {str(e)}"


@agent.tool(sequential=True)
async def edit_field(
    ctx: RunContext[TemplateAgentContext],
    tag: str,
    new_tag: str | None = None,
    new_alias: str | None = None,
    new_value: str | None = None,
    sync_new_tag_changes: bool = True,
) -> str:
    """
    Edit an existing template field's properties.

    By default, changing the tag will also update the alias and value to match.
    Set sync_new_tag_changes=False to provide explicit values or change only the tag.

    Args:
        tag: The current tag of the field to edit (e.g., "declarant_name")
        new_tag: New tag value (e.g., "declarant.name"). None to keep existing.
        new_alias: New display name. Requires sync_new_tag_changes=False if provided with new_tag.
        new_value: New text value. Requires sync_new_tag_changes=False if provided with new_tag.
        sync_new_tag_changes: If True (default), changing tag auto-generates alias/value.
                              MUST be False if providing explicit new_alias or new_value with new_tag.

    Examples:
        # Rename field completely (tag, alias, value all sync)
        edit_field(tag="client_name", new_tag="client.name")

        # Rename with custom alias/value - must set sync_new_tag_changes=False
        edit_field(tag="client_name", new_tag="client.name", new_alias="Full Name", sync_new_tag_changes=False)

        # Change only the tag, keep existing alias/value
        edit_field(tag="client_name", new_tag="client.name", sync_new_tag_changes=False)

        # Change just alias or value (no new_tag, sync doesn't apply)
        edit_field(tag="client_name", new_alias="Full Name")

    Returns:
        Status message indicating success or failure
    """
    if ctx.deps.document is None:
        return "No document loaded. Use load_document first."

    if new_tag is None and new_alias is None and new_value is None:
        return "No changes specified. Provide at least one of: new_tag, new_alias, or new_value."

    # Validation: explicit values with sync_new_tag_changes=True is ambiguous
    if sync_new_tag_changes and new_tag is not None:
        if new_alias is not None or new_value is not None:
            provided = []
            if new_alias is not None:
                provided.append("new_alias")
            if new_value is not None:
                provided.append("new_value")
            return (
                f"Ambiguous call: {', '.join(provided)} provided with sync_new_tag_changes=True. "
                f"Set sync_new_tag_changes=False when providing explicit values with new_tag."
            )
        # Auto-generate alias and value
        new_alias = new_tag.replace(".", " ").replace("_", " ").title()
        new_value = f"[{new_tag}]"

    try:
        # Get current info for logging
        working_path = ctx.deps.get_working_path()
        if working_path and working_path.exists():
            controls = read_content_controls_detailed(str(working_path))
        else:
            controls = read_content_controls_detailed(str(ctx.deps.get_document_path()))

        # Check if field exists
        field_info = None
        for ctrl in controls:
            if ctrl["tag"] == tag:
                field_info = ctrl
                break

        if field_info is None:
            available_tags = [ctrl["tag"] for ctrl in controls]
            return f"Field with tag '{tag}' not found. Available tags: {available_tags}"

        # Update the field
        count = update_content_control(
            ctx.deps.document,
            tag=tag,
            new_tag=new_tag,
            new_alias=new_alias,
            new_value=new_value,
        )

        if count > 0:
            # Build change summary
            changes = []
            if new_tag is not None:
                changes.append(f"tag: '{tag}' → '{new_tag}'")
            if new_alias is not None:
                changes.append(f"display name: '{field_info['alias']}' → '{new_alias}'")
            if new_value is not None:
                old_val = field_info['value'][:30] + "..." if len(field_info['value']) > 30 else field_info['value']
                new_val = new_value[:30] + "..." if len(new_value) > 30 else new_value
                changes.append(f"value: '{old_val}' → '{new_val}'")

            change_summary = ", ".join(changes)

            # Save working copy
            if working_path:
                ctx.deps.document.save(str(working_path))

            logfire.info(f"Edited field {tag}", new_tag=new_tag, new_alias=new_alias, new_value=new_value)
            return f"Edited field '{tag}': {change_summary}"
        else:
            return f"Failed to edit field '{tag}'"

    except Exception as e:
        logfire.error(f"Failed to edit field: {e}")
        return f"Error editing field: {str(e)}"


@agent.tool(sequential=True)
async def split_field(
    ctx: RunContext[TemplateAgentContext],
    tag: str,
    new_tags: list[str],
) -> str:
    """
    Split one field into multiple adjacent fields.

    Replaces a single field with multiple fields. Each new field gets a placeholder
    value of "[tag_name]". Useful for splitting a name field into first/middle/last.

    Args:
        tag: The tag of the existing field to split (e.g., "declarant.name")
        new_tags: List of new tag names for the split fields
                  (e.g., ["declarant.first_name", "declarant.middle_name", "declarant.last_name"])

    Example:
        split_field(tag="declarant.name", new_tags=["declarant.first_name", "declarant.middle_name", "declarant.last_name"])
        → Replaces [declarant.name] with [declarant.first_name] [declarant.middle_name] [declarant.last_name]

    Returns:
        Status message indicating success or failure
    """
    if ctx.deps.document is None:
        return "No document loaded. Use load_document first."

    if not new_tags or len(new_tags) < 2:
        return "Must provide at least 2 new tags to split a field."

    try:
        # Check if field exists
        working_path = ctx.deps.get_working_path()
        if working_path and working_path.exists():
            controls = read_content_controls_detailed(str(working_path))
        else:
            controls = read_content_controls_detailed(str(ctx.deps.get_document_path()))

        field_info = None
        for ctrl in controls:
            if ctrl["tag"] == tag:
                field_info = ctrl
                break

        if field_info is None:
            available_tags = [ctrl["tag"] for ctrl in controls]
            return f"Field with tag '{tag}' not found. Available tags: {available_tags}"

        # Split the field
        count = split_content_control(ctx.deps.document, tag, new_tags)

        if count > 0:
            new_tags_str = ", ".join([f"[{t}]" for t in new_tags])

            # Save working copy
            if working_path:
                ctx.deps.document.save(str(working_path))

            logfire.info(f"Split field {tag} into {count} fields", new_tags=new_tags)
            return f"Split [{tag}] into {count} fields: {new_tags_str}"
        else:
            return f"Failed to split field '{tag}'"

    except Exception as e:
        logfire.error(f"Failed to split field: {e}")
        return f"Error splitting field: {str(e)}"


@agent.tool(sequential=True)
async def reset_working_copy(ctx: RunContext[TemplateAgentContext]) -> str:
    """
    Reset the working copy by deleting it and reloading from the original document.

    Use this to discard all modifications and start fresh.

    Returns:
        Status message indicating success
    """
    if ctx.deps.document_filename is None:
        return "No document loaded."

    try:
        # Delete working copy if it exists
        working_path = ctx.deps.get_working_path()
        if working_path and working_path.exists():
            working_path.unlink()
            logfire.info(f"Deleted working copy: {working_path}")

        # Reload from original
        ctx.deps.document = None
        success, message = ctx.deps.load_document()

        if success:
            return f"Working copy reset. Reloaded original document: {ctx.deps.document_filename}"
        else:
            return f"Reset working copy but failed to reload: {message}"

    except Exception as e:
        logfire.error(f"Failed to reset working copy: {e}")
        return f"Error resetting working copy: {str(e)}"


def _format_fields(controls: list[dict]) -> str:
    """Format template fields for display"""
    result = ""
    for i, ctrl in enumerate(controls, 1):
        result += f"{i}. Tag: {ctrl['tag']}\n"
        result += f"   Display Name: {ctrl['alias']}\n"
        result += f"   Current Value: {ctrl['value'][:50]}{'...' if len(ctrl['value']) > 50 else ''}\n\n"
    return result
