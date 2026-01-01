from docx import Document
from docx.oxml.ns import qn
import os 

os.chdir(os.path.dirname(os.path.abspath(__file__)))
print(os.getcwd())

W_SDT = qn("w:sdt")
W_SDT_PR = qn("w:sdtPr")
W_TAG = qn("w:tag")
W_VAL = qn("w:val")
W_SDT_CONTENT = qn("w:sdtContent")
W_T = qn("w:t")

def read_content_controls(path: str) -> dict[str, str]:
    doc = Document(path)
    result = {}

    for sdt in doc._element.iter(W_SDT):
        sdtPr = sdt.find(W_SDT_PR)
        if sdtPr is None:
            continue

        tag = sdtPr.find(W_TAG)
        if tag is None:
            continue

        key = tag.get(W_VAL)

        texts = [
            t.text for t in sdt.find(W_SDT_CONTENT).iter(W_T)
            if t.text
        ]
        result[key] = "".join(texts)

    return result




def set_content_control(path_in: str, path_out: str, key: str, value: str) -> None:
    doc = Document(path_in)

    for sdt in doc._element.iter(W_SDT):
        sdtPr = sdt.find(W_SDT_PR)
        if sdtPr is None:
            continue

        tag = sdtPr.find(W_TAG)
        if tag is None or tag.get(W_VAL) != key:
            continue

        content = sdt.find(W_SDT_CONTENT)

        # Clear all existing text nodes
        for t in content.iter(W_T):
            t.text = ""

        # Write into the first text node (common case)
        first = next(content.iter(W_T), None)
        if first is None:
            raise RuntimeError(f"Content control '{key}' has no text node")

        first.text = value
        break

    doc.save(path_out)

# path = "../test_template.docx"
path = "simple legal will template.docx"
values = read_content_controls(path)
print(values)

set_content_control(
    path_in=path,
    path_out=path.replace(".docx", "_filled.docx"),
    key="field1tag",
    value="Acme Corp"
)

values = read_content_controls(path)
print(values)

# interactive explore the doc (to be done in a notebook):
doc = Document(path)

print(f"\n{'='*20} DOCUMENT STRUCTURE {'='*20}")

print(f"\n[PARAGRAPHS] Found {len(doc.paragraphs)} paragraphs")
for i, p in enumerate(doc.paragraphs):
    if not p.text.strip(): 
        continue # skip empty
    print(f"  P{i}: '{p.text[:40]}...' (Style: {p.style.name})")
    # specific formatting is in runs
    for j, run in enumerate(p.runs):
        if run.text.strip():
            print(f"    - Run {j}: '{run.text}' (Bold: {run.bold}, Italic: {run.italic})")

print(f"\n[TABLES] Found {len(doc.tables)} tables")
for i, table in enumerate(doc.tables):
    print(f"  Table {i}: {len(table.rows)} rows x {len(table.columns)} cols")
    # Quick preview of first row
    if table.rows:
        row_data = [cell.text for cell in table.rows[0].cells]
        print(f"    Row 0: {row_data}")

# For your specific use case (Content Controls / SDT), they are often hidden in XML
# doc.element.xml contains the full underlying XML tree
print(f"\n[XML INSPECTION] Root element tag: {doc.element.tag}")

import lxml.etree
xml_str = lxml.etree.tostring(doc.element, pretty_print=True, encoding='unicode')

xml_path = path.replace(".docx", ".xml")
with open(xml_path, "w") as f:
    f.write(xml_str)

print(xml_str)

print(f"\n{'='*20} RECURSIVE TEXT SEARCH {'='*20}")
# Helper to find ALL text, even inside Content Controls (w:sdt)
def iter_text(element):
    for child in element:
        # If it's a text node (w:t), yield it
        if child.tag == qn('w:t'):
            if child.text:
                yield child.text
        # Recurse into children
        yield from iter_text(child)

all_text = list(iter_text(doc.element))
print(f"Found {len(all_text)} text fragments:")
print(all_text)