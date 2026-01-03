"""
DOCX Content Controls Utilities

Functions for reading, creating, and modifying Word content controls (SDT elements).
Content controls are fillable fields in Word documents that can be programmatically
manipulated.

Usage:
    from docx_content_controls import (
        read_content_controls,
        set_content_control,
        add_content_control_paragraph,
        add_inline_content_control,
    )
"""

from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import random
import lxml.etree


# =============================================================================
# XML Namespace Constants
# =============================================================================

W_SDT = qn("w:sdt")
W_SDT_PR = qn("w:sdtPr")
W_SDT_END_PR = qn("w:sdtEndPr")
W_SDT_CONTENT = qn("w:sdtContent")
W_TAG = qn("w:tag")
W_ALIAS = qn("w:alias")
W_ID = qn("w:id")
W_VAL = qn("w:val")
W_T = qn("w:t")
W_R = qn("w:r")
W_P = qn("w:p")
W_PLACEHOLDER = qn("w:placeholder")
W_DOC_PART = qn("w:docPart")


# =============================================================================
# Low-Level Element Creation
# =============================================================================

def _make_element(tag: str, **attribs) -> OxmlElement:
    """Create an OxmlElement with optional w: namespaced attributes."""
    el = OxmlElement(tag)
    for key, val in attribs.items():
        el.set(qn(f"w:{key}"), str(val))
    return el


def create_content_control_element(
    tag: str,
    value: str,
    alias: str | None = None,
    sdt_id: int | None = None,
    block_level: bool = False,
) -> OxmlElement:
    """
    Create a new content control (w:sdt) XML element.

    Args:
        tag: The programmatic tag/key for the content control
        value: The initial text value
        alias: Display name shown in Word (defaults to tag)
        sdt_id: Unique ID (auto-generated if not provided)
        block_level: If True, creates block-level SDT containing a paragraph.
                     If False, creates inline SDT containing just a run.

    Returns:
        OxmlElement: The constructed w:sdt element
    """
    if alias is None:
        alias = tag
    if sdt_id is None:
        sdt_id = random.randint(100000000, 999999999)

    # Build w:sdt
    sdt = _make_element("w:sdt")

    # Build w:sdtPr (properties)
    sdt_pr = _make_element("w:sdtPr")
    sdt_pr.append(_make_element("w:alias", val=alias))
    sdt_pr.append(_make_element("w:tag", val=tag))
    sdt_pr.append(_make_element("w:id", val=sdt_id))

    # Add placeholder
    placeholder = _make_element("w:placeholder")
    placeholder.append(_make_element("w:docPart", val="DefaultPlaceholder_-1854013440"))
    sdt_pr.append(placeholder)

    sdt.append(sdt_pr)

    # Build w:sdtEndPr (usually empty)
    sdt.append(_make_element("w:sdtEndPr"))

    # Build w:sdtContent
    sdt_content = _make_element("w:sdtContent")

    # Create the text run
    run = _make_element("w:r")
    text = _make_element("w:t")
    text.text = value
    run.append(text)

    if block_level:
        # Block-level: wrap run in a paragraph
        para = _make_element("w:p")
        para.append(run)
        sdt_content.append(para)
    else:
        # Inline: just the run
        sdt_content.append(run)

    sdt.append(sdt_content)

    return sdt


# =============================================================================
# Document Modification Functions
# =============================================================================

def add_content_control_paragraph(
    doc: Document,
    tag: str,
    value: str,
    alias: str | None = None,
    position: int | None = None,
) -> OxmlElement:
    """
    Add a new block-level content control as a standalone paragraph.

    Args:
        doc: The python-docx Document object
        tag: The programmatic tag/key for the content control
        value: The initial text value
        alias: Display name (defaults to tag)
        position: Index to insert at (appends to end if None)

    Returns:
        The created w:sdt element
    """
    sdt = create_content_control_element(tag, value, alias, block_level=True)
    body = doc._element.body

    if position is None:
        # Append before sectPr if it exists, otherwise at end
        sect_pr = body.find(qn("w:sectPr"))
        if sect_pr is not None:
            sect_pr.addprevious(sdt)
        else:
            body.append(sdt)
    else:
        children = list(body)
        if position < len(children):
            children[position].addprevious(sdt)
        else:
            body.append(sdt)

    return sdt


def add_inline_content_control(
    paragraph,
    tag: str,
    value: str,
    alias: str | None = None,
    position: int | None = None,
) -> OxmlElement:
    """
    Add an inline content control within an existing paragraph.

    Args:
        paragraph: A python-docx Paragraph object
        tag: The programmatic tag/key for the content control
        value: The initial text value
        alias: Display name (defaults to tag)
        position: Index within paragraph's children (appends if None)

    Returns:
        The created w:sdt element
    """
    sdt = create_content_control_element(tag, value, alias, block_level=False)
    p_element = paragraph._element

    if position is None:
        p_element.append(sdt)
    else:
        children = list(p_element)
        if position < len(children):
            children[position].addprevious(sdt)
        else:
            p_element.append(sdt)

    return sdt


def _collect_all_paragraphs(doc: Document) -> list:
    """
    Collect all paragraphs from the entire document, including:
    - Body paragraphs
    - Table cells
    - Headers and footers
    """
    paragraphs = []

    # Body paragraphs
    paragraphs.extend(doc.paragraphs)

    # Table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)

    # Headers and footers
    for section in doc.sections:
        try:
            if section.header:
                paragraphs.extend(section.header.paragraphs)
        except Exception:
            pass
        try:
            if section.footer:
                paragraphs.extend(section.footer.paragraphs)
        except Exception:
            pass

    return paragraphs


def wrap_text_in_content_control(
    doc: Document,
    search_text: str,
    tag: str,
    alias: str | None = None,
    first_only: bool = False,
) -> int:
    """
    Find text in the document and wrap it in a content control.

    Searches the ENTIRE document including body text, tables, headers, and footers.
    Handles text that spans multiple runs (common when Word applies formatting).

    Args:
        doc: The python-docx Document object
        search_text: The exact text to find and wrap
        tag: The programmatic tag/key for the content control
        alias: Display name (defaults to tag)
        first_only: If True, only wrap the first occurrence. If False, wrap all occurrences.

    Returns:
        Number of content controls created
    """
    count = 0

    # Search all paragraphs in the document (body, tables, headers, footers)
    for paragraph in _collect_all_paragraphs(doc):
        # Build concatenated text and track run boundaries
        runs = paragraph.runs
        if not runs:
            continue

        combined_text = ""
        run_boundaries = []  # [(start_pos, end_pos, run_index), ...]

        for i, run in enumerate(runs):
            start = len(combined_text)
            combined_text += run.text or ""
            end = len(combined_text)
            run_boundaries.append((start, end, i))

        # Search for text in combined paragraph text
        search_start = 0
        while True:
            idx = combined_text.find(search_text, search_start)
            if idx == -1:
                break

            match_end = idx + len(search_text)

            # Find which runs contain the match
            first_run_idx = None
            last_run_idx = None
            for start, end, run_idx in run_boundaries:
                if start < match_end and end > idx:
                    if first_run_idx is None:
                        first_run_idx = run_idx
                    last_run_idx = run_idx

            if first_run_idx is None:
                search_start = idx + 1
                continue

            # Calculate offsets within first and last runs
            first_run_start, _, _ = run_boundaries[first_run_idx]
            last_run_start, _, _ = run_boundaries[last_run_idx]

            offset_in_first = idx - first_run_start
            offset_in_last = match_end - last_run_start

            # Get the parent element and run positions
            parent = runs[first_run_idx]._element.getparent()
            first_elem_idx = list(parent).index(runs[first_run_idx]._element)

            # Collect text parts
            first_run_text = runs[first_run_idx].text or ""
            last_run_text = runs[last_run_idx].text or ""

            before_text = first_run_text[:offset_in_first]
            after_text = last_run_text[offset_in_last:]

            # Remove runs that are part of the match (in reverse order to preserve indices)
            for i in range(last_run_idx, first_run_idx - 1, -1):
                parent.remove(runs[i]._element)

            # Insert replacement elements
            insert_pos = first_elem_idx

            # Add "before" text if any
            if before_text:
                before_run = _make_element("w:r")
                before_t = _make_element("w:t")
                before_t.text = before_text
                if before_text.endswith(" "):
                    before_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                before_run.append(before_t)
                parent.insert(insert_pos, before_run)
                insert_pos += 1

            # Add the content control
            sdt = create_content_control_element(
                tag=tag,
                value=search_text,
                alias=alias,
                block_level=False,
            )
            parent.insert(insert_pos, sdt)
            insert_pos += 1

            # Add "after" text if any
            if after_text:
                after_run = _make_element("w:r")
                after_t = _make_element("w:t")
                after_t.text = after_text
                if after_text.startswith(" "):
                    after_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                after_run.append(after_t)
                parent.insert(insert_pos, after_run)

            count += 1
            if first_only:
                return count

            # Paragraph structure changed, break and let outer loop re-process if needed
            break

    return count


def set_content_control_value(
    doc: Document,
    tag: str,
    value: str,
) -> bool:
    """
    Update all content controls with the given tag in a Document object.

    Args:
        doc: The python-docx Document object
        tag: The tag/key of the content control(s) to update
        value: The new value to set

    Returns:
        True if at least one was found and updated, False if none found
    """
    found = False
    for sdt in doc._element.iter(W_SDT):
        sdtPr = sdt.find(W_SDT_PR)
        if sdtPr is None:
            continue

        tag_el = sdtPr.find(W_TAG)
        if tag_el is None or tag_el.get(W_VAL) != tag:
            continue

        content = sdt.find(W_SDT_CONTENT)

        # Clear all existing text nodes
        for t in content.iter(W_T):
            t.text = ""

        # Write into the first text node
        first = next(content.iter(W_T), None)
        if first is None:
            raise RuntimeError(f"Content control '{tag}' has no text node")

        first.text = value
        found = True

    return found


def delete_content_control(
    doc: Document,
    tag: str,
    preserve_text: bool = True,
) -> int:
    """
    Delete content control(s) with the given tag from a Document.

    Args:
        doc: The python-docx Document object
        tag: The tag of the content control(s) to delete
        preserve_text: If True, keep the text content in place (unwrap).
                       If False, remove both the control and its text.

    Returns:
        Number of content controls deleted
    """
    count = 0
    # Collect SDTs to delete (can't modify while iterating)
    sdts_to_delete = []

    for sdt in doc._element.iter(W_SDT):
        sdtPr = sdt.find(W_SDT_PR)
        if sdtPr is None:
            continue

        tag_el = sdtPr.find(W_TAG)
        if tag_el is None or tag_el.get(W_VAL) != tag:
            continue

        sdts_to_delete.append(sdt)

    for sdt in sdts_to_delete:
        parent = sdt.getparent()
        if parent is None:
            continue

        if preserve_text:
            # Get the content and insert it in place of the SDT
            content = sdt.find(W_SDT_CONTENT)
            if content is not None:
                # Get position of SDT in parent
                sdt_index = list(parent).index(sdt)
                # Remove the SDT first
                parent.remove(sdt)
                # Insert all children of sdtContent at the original position
                for i, child in enumerate(list(content)):
                    parent.insert(sdt_index + i, child)
            else:
                parent.remove(sdt)
        else:
            # Just remove the entire SDT
            parent.remove(sdt)

        count += 1

    return count


def split_content_control(
    doc: Document,
    tag: str,
    new_tags: list[str],
) -> int:
    """
    Split a content control into multiple adjacent content controls.

    Replaces one field with N fields. Each new field gets a placeholder value
    of "[tag_name]". Useful for splitting a name field into first/middle/last.

    Args:
        doc: The python-docx Document object
        tag: The tag of the content control to split
        new_tags: List of new tag names (e.g., ["person.first", "person.middle", "person.last"])

    Returns:
        Number of new content controls created (0 if original not found)
    """
    if not new_tags:
        return 0

    # Find the content control to split
    target_sdt = None
    for sdt in doc._element.iter(W_SDT):
        sdtPr = sdt.find(W_SDT_PR)
        if sdtPr is None:
            continue
        tag_el = sdtPr.find(W_TAG)
        if tag_el is not None and tag_el.get(W_VAL) == tag:
            target_sdt = sdt
            break

    if target_sdt is None:
        return 0

    parent = target_sdt.getparent()
    if parent is None:
        return 0

    # Get position of original SDT
    sdt_index = list(parent).index(target_sdt)

    # Remove the original SDT
    parent.remove(target_sdt)

    # Insert new content controls with spaces between them
    insert_pos = sdt_index
    for i, new_tag in enumerate(new_tags):
        # Generate alias from tag: "person.first_name" -> "Person First Name"
        alias = new_tag.replace(".", " ").replace("_", " ").title()
        placeholder = f"[{new_tag}]"

        # Create the new content control
        new_sdt = create_content_control_element(
            tag=new_tag,
            value=placeholder,
            alias=alias,
            block_level=False,
        )
        parent.insert(insert_pos, new_sdt)
        insert_pos += 1

        # Add a space run between fields (except after the last one)
        if i < len(new_tags) - 1:
            space_run = _make_element("w:r")
            space_t = _make_element("w:t")
            space_t.text = " "
            space_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            space_run.append(space_t)
            parent.insert(insert_pos, space_run)
            insert_pos += 1

    return len(new_tags)


def update_content_control(
    doc: Document,
    tag: str,
    new_tag: str | None = None,
    new_alias: str | None = None,
    new_value: str | None = None,
) -> int:
    """
    Update properties of content control(s) with the given tag.

    Args:
        doc: The python-docx Document object
        tag: The tag of the content control(s) to update
        new_tag: New tag value (None to keep existing)
        new_alias: New alias/display name (None to keep existing)
        new_value: New text value (None to keep existing)

    Returns:
        Number of content controls updated
    """
    count = 0

    for sdt in doc._element.iter(W_SDT):
        sdtPr = sdt.find(W_SDT_PR)
        if sdtPr is None:
            continue

        tag_el = sdtPr.find(W_TAG)
        if tag_el is None or tag_el.get(W_VAL) != tag:
            continue

        # Update tag if requested
        if new_tag is not None:
            tag_el.set(W_VAL, new_tag)

        # Update alias if requested
        if new_alias is not None:
            alias_el = sdtPr.find(W_ALIAS)
            if alias_el is not None:
                alias_el.set(W_VAL, new_alias)
            else:
                # Create alias element if it doesn't exist
                alias_el = _make_element("w:alias", val=new_alias)
                # Insert after tag element or at beginning
                tag_idx = list(sdtPr).index(tag_el)
                sdtPr.insert(tag_idx, alias_el)

        # Update value if requested
        if new_value is not None:
            content = sdt.find(W_SDT_CONTENT)
            if content is not None:
                # Clear all existing text nodes
                for t in content.iter(W_T):
                    t.text = ""
                # Write into the first text node
                first = next(content.iter(W_T), None)
                if first is not None:
                    first.text = new_value

        count += 1

    return count


# =============================================================================
# Reading Functions
# =============================================================================

def read_content_controls(path: str | Path) -> dict[str, str]:
    """
    Read content controls and return a simple tag -> value mapping.

    Args:
        path: Path to the DOCX file

    Returns:
        Dict mapping content control tags to their text values
    """
    doc = Document(str(path))
    result = {}

    for sdt in doc._element.iter(W_SDT):
        sdtPr = sdt.find(W_SDT_PR)
        if sdtPr is None:
            continue

        tag = sdtPr.find(W_TAG)
        if tag is None:
            continue

        key = tag.get(W_VAL)
        content = sdt.find(W_SDT_CONTENT)
        texts = [t.text for t in content.iter(W_T) if t.text] if content is not None else []
        result[key] = "".join(texts)

    return result


def read_content_controls_detailed(path: str | Path) -> list[dict]:
    """
    Read content controls with full metadata.

    Args:
        path: Path to the DOCX file

    Returns:
        List of dicts with: tag, alias, value, id
    """
    doc = Document(str(path))
    result = []

    for sdt in doc._element.iter(W_SDT):
        sdtPr = sdt.find(W_SDT_PR)
        if sdtPr is None:
            continue

        tag_el = sdtPr.find(W_TAG)
        if tag_el is None:
            continue

        tag = tag_el.get(W_VAL)

        alias_el = sdtPr.find(W_ALIAS)
        alias = alias_el.get(W_VAL) if alias_el is not None else tag

        id_el = sdtPr.find(W_ID)
        sdt_id = id_el.get(W_VAL) if id_el is not None else None

        content = sdt.find(W_SDT_CONTENT)
        texts = [t.text for t in content.iter(W_T) if t.text] if content is not None else []
        value = "".join(texts)

        result.append({
            "tag": tag,
            "alias": alias,
            "value": value,
            "id": sdt_id
        })

    return result


# =============================================================================
# File I/O Functions
# =============================================================================

def export_to_xml(doc: Document, path: str | Path) -> Path:
    """
    Export a Document's XML to a file.

    Args:
        doc: The python-docx Document object
        path: Output path for the XML file

    Returns:
        Path to the exported XML file
    """
    path = Path(path)
    xml_str = lxml.etree.tostring(doc._element, pretty_print=True, encoding='unicode')
    path.write_text(xml_str)
    return path


def save_document(doc: Document, base_path: str | Path, suffix: str = "") -> tuple[Path, Path]:
    """
    Save a document as both DOCX and XML files.

    Args:
        doc: The python-docx Document object
        base_path: Base path without extension (e.g., "my_document")
        suffix: Optional suffix to add (e.g., "_modified")

    Returns:
        Tuple of (docx_path, xml_path)
    """
    base_path = Path(base_path)
    stem = base_path.stem + suffix
    parent = base_path.parent

    docx_path = parent / f"{stem}.docx"
    xml_path = parent / f"{stem}.xml"

    doc.save(str(docx_path))
    export_to_xml(doc, xml_path)

    return docx_path, xml_path


def set_content_control(path_in: str | Path, path_out: str | Path, key: str, value: str) -> Document:
    """
    Load a document, update a content control, and save to a new file.

    Args:
        path_in: Input DOCX file path
        path_out: Output DOCX file path
        key: The tag of the content control to update
        value: The new value to set

    Returns:
        The modified Document object
    """
    doc = Document(str(path_in))

    if not set_content_control_value(doc, key, value):
        raise ValueError(f"Content control with tag '{key}' not found")

    doc.save(str(path_out))
    return doc


# =============================================================================
# Demo Document Generation
# =============================================================================

def generate_template_document() -> Document:
    """
    Generate a simple template document with a few pre-existing content controls.
    This simulates a template that would be filled in later.

    Returns:
        A python-docx Document with minimal content controls
    """
    doc = Document()
    doc.add_heading("Legal Document Template", 0)

    # Pre-existing content controls section
    doc.add_heading("Pre-existing Content Controls", level=1)
    doc.add_paragraph("These content controls already exist in the template:")
    doc.add_paragraph("")

    # Block-level: Client name
    doc.add_paragraph("Client Name:")
    add_content_control_paragraph(doc, tag="client_name", value="[Enter Name]", alias="Client Name")

    doc.add_paragraph("")

    # Inline: Declaration with city/state
    p = doc.add_paragraph("Client resides in ")
    add_inline_content_control(p, tag="client_city", value="[City]", alias="City")
    p.add_run(", ")
    add_inline_content_control(p, tag="client_state", value="[State]", alias="State")
    p.add_run(".")

    # Empty sections for later additions
    doc.add_paragraph("")
    doc.add_heading("Additional Information", level=1)
    doc.add_paragraph("This section will have content controls added programmatically.")

    doc.add_paragraph("")
    doc.add_heading("Legal Terms", level=1)
    doc.add_paragraph("The total value is $100,000 as agreed upon by both parties.")

    doc.add_paragraph("")
    doc.add_heading("Signature", level=1)
    doc.add_paragraph("_" * 40)
    doc.add_paragraph("(Signature will be added here)")

    return doc


def run_demo(output_dir: str | Path | None = None) -> dict:
    """
    Run a full demo showing the workflow:
    1. Generate a template with a few pre-existing content controls
    2. Save as DOCX and XML
    3. Load, modify existing controls, and add new ones using each function
    4. Save the modified version

    Args:
        output_dir: Directory for output files (defaults to script directory)

    Returns:
        Dict with paths to generated files and content control values
    """
    if output_dir is None:
        output_dir = Path(__file__).parent
    output_dir = Path(output_dir)

    results = {
        "files": {},
        "original_values": {},
        "modified_values": {},
    }

    print("\n" + "=" * 60)
    print("STEP 1: Generate template document (few pre-existing controls)")
    print("=" * 60)

    doc = generate_template_document()
    base_path = output_dir / "content_control_demo"
    docx_path, xml_path = save_document(doc, base_path)
    results["files"]["original_docx"] = str(docx_path)
    results["files"]["original_xml"] = str(xml_path)
    print(f"  Saved: {docx_path}")
    print(f"  Saved: {xml_path}")

    print("\n" + "=" * 60)
    print("STEP 2: Read pre-existing content controls")
    print("=" * 60)

    values = read_content_controls(docx_path)
    results["original_values"] = values
    print(f"  Found {len(values)} pre-existing content controls:")
    for tag, value in values.items():
        print(f"    {tag}: '{value}'")

    print("\n" + "=" * 60)
    print("STEP 3: Modify and add content controls")
    print("=" * 60)

    # Reload document for modifications
    doc = Document(str(docx_path))

    # 3a. Modify ONE existing control using set_content_control_value()
    print("\n  Using set_content_control_value():")
    if set_content_control_value(doc, "client_name", "John Smith"):
        print("    Updated 'client_name': '[Enter Name]' -> 'John Smith'")

    # 3b. Add block-level content control using add_content_control_paragraph()
    print("\n  Using add_content_control_paragraph():")
    add_content_control_paragraph(
        doc,
        tag="additional_info",
        value="Additional details go here",
        alias="Additional Info"
    )
    print("    Added block-level control: 'additional_info'")

    # 3c. Add inline content controls to existing paragraph
    print("\n  Using add_inline_content_control():")
    # Find the "Additional Information" section and add to it
    for para in doc.paragraphs:
        if "content controls added programmatically" in para.text:
            para.clear()
            para.add_run("Contact: ")
            add_inline_content_control(para, tag="contact_email", value="email@example.com", alias="Email")
            para.add_run(" | Phone: ")
            add_inline_content_control(para, tag="contact_phone", value="555-1234", alias="Phone")
            print("    Added inline controls: 'contact_email', 'contact_phone'")
            break

    # 3d. Use wrap_text_in_content_control() to wrap existing text
    print("\n  Using wrap_text_in_content_control():")
    wrapped = wrap_text_in_content_control(
        doc,
        search_text="$100,000",
        tag="total_value",
        alias="Total Value"
    )
    if wrapped:
        print(f"    Wrapped existing text '$100,000' -> control 'total_value'")

    # 3e. Add signature block-level control
    print("\n  Using add_content_control_paragraph() for signature:")
    # Find and replace signature placeholder
    for para in doc.paragraphs:
        if "(Signature will be added here)" in para.text:
            para.clear()
            break
    add_content_control_paragraph(doc, tag="signature", value="John Smith", alias="Signature")
    print("    Added block-level control: 'signature'")

    print("\n" + "=" * 60)
    print("STEP 4: Save modified document")
    print("=" * 60)

    docx_path_mod, xml_path_mod = save_document(doc, base_path, suffix="_modified")
    results["files"]["modified_docx"] = str(docx_path_mod)
    results["files"]["modified_xml"] = str(xml_path_mod)
    print(f"  Saved: {docx_path_mod}")
    print(f"  Saved: {xml_path_mod}")

    print("\n" + "=" * 60)
    print("STEP 5: Verify - compare original vs modified")
    print("=" * 60)

    modified_values = read_content_controls(docx_path_mod)
    results["modified_values"] = modified_values

    print(f"\n  Original ({len(values)} controls):")
    for tag, value in values.items():
        print(f"    {tag}: '{value}'")

    print(f"\n  Modified ({len(modified_values)} controls):")
    for tag, value in modified_values.items():
        original = values.get(tag)
        if original is None:
            print(f"    {tag}: '{value}' [NEW]")
        elif original != value:
            print(f"    {tag}: '{value}' [was: '{original}']")
        else:
            print(f"    {tag}: '{value}'")

    print("\n" + "=" * 60)
    print("SUMMARY")
    print("=" * 60)
    print(f"""
Generated files:
  - {results['files']['original_docx']} (template with {len(values)} controls)
  - {results['files']['original_xml']}
  - {results['files']['modified_docx']} (filled with {len(modified_values)} controls)
  - {results['files']['modified_xml']}

Functions demonstrated:
  - generate_template_document() -> create template with few controls
  - save_document() -> save as both DOCX and XML
  - read_content_controls() -> read tag -> value mapping
  - set_content_control_value() -> modify existing control
  - add_content_control_paragraph() -> add new block-level control
  - add_inline_content_control() -> add inline control to paragraph
  - wrap_text_in_content_control() -> wrap existing text in control
""")

    return results


# =============================================================================
# Main Entry Point
# =============================================================================

if __name__ == "__main__":
    import os
    os.chdir(Path(__file__).parent)
    run_demo(output_dir="documents")
